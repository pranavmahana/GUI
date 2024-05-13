import tkinter as tk
from tkinter import filedialog, ttk, messagebox as messagebox
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from scipy.signal import welch
from docx import Document
from docx.shared import Inches


class VibrationAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("Vibration Analyzer")

        # Variables to store user input
        self.sensitivity = tk.DoubleVar(value=10)
        self.sampling_frequency = tk.DoubleVar(value=25000)
        self.csv_file_path = tk.StringVar()
        self.velocity_csv_file_path = tk.StringVar()  

        # Main notebook for tabs
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(expand=True, fill="both")

        # Input tab
        self.input_frame = tk.Frame(self.notebook)
        self.notebook.add(self.input_frame, text="Input")

        # Sensitivity input
        tk.Label(self.input_frame, text="Sensitivity:").grid(row=0, column=0, sticky="w")
        tk.Entry(self.input_frame, textvariable=self.sensitivity).grid(row=0, column=1, sticky="w")

        # Sampling frequency input
        tk.Label(self.input_frame, text="Sampling Frequency:").grid(row=1, column=0, sticky="w")
        tk.Entry(self.input_frame, textvariable=self.sampling_frequency).grid(row=1, column=1, sticky="w")

        # Load CSV button for main data
        tk.Label(self.input_frame, text="Main Data CSV/File:").grid(row=2, column=0, sticky="w")
        self.csv_entry = tk.Entry(self.input_frame, textvariable=self.csv_file_path, state="readonly")
        self.csv_entry.grid(row=2, column=1, sticky="we", columnspan=2)
        tk.Button(self.input_frame, text="Load Main Data CSV/File", command=lambda: self.load_file(self.csv_file_path)).grid(row=3, column=0, columnspan=2)

        # Load CSV button for velocity data
        tk.Label(self.input_frame, text="Velocity CSV/File:").grid(row=4, column=0, sticky="w")
        self.velocity_csv_entry = tk.Entry(self.input_frame, textvariable=self.velocity_csv_file_path, state="readonly")
        self.velocity_csv_entry.grid(row=4, column=1, sticky="we", columnspan=2)
        tk.Button(self.input_frame, text="Load Velocity CSV/File", command=lambda: self.load_file(self.velocity_csv_file_path)).grid(row=5, column=0, columnspan=2)

        # Plot G-levels button
        tk.Button(self.input_frame, text="Plot G-levels", command=self.plot_glevels).grid(row=6, column=0, columnspan=2)

        # Plot PSD button
        tk.Button(self.input_frame, text="Plot PSD", command=self.plot_psd).grid(row=7, column=0, columnspan=2)

        # Save button
        self.save_button = tk.Button(self.input_frame, text="Save Images", command=self.save_images)

        # G-level plots tab
        self.glevel_plots_frame = tk.Frame(self.notebook)
        self.notebook.add(self.glevel_plots_frame, text="G-level Plots")

        # Canvas for plotting G-levels
        self.glevel_fig = plt.Figure(figsize=(14, 10))
        self.glevel_canvas = FigureCanvasTkAgg(self.glevel_fig, self.glevel_plots_frame)
        self.glevel_toolbar = NavigationToolbar2Tk(self.glevel_canvas, self.glevel_plots_frame)
        self.glevel_toolbar.update()
        self.glevel_canvas.get_tk_widget().pack(side="top", fill="both", expand=True)

        # Toolbar for G-level plots
        self.glevel_toolbar.pack(side="bottom", fill="x")

        # Plot click event for G-level plots
        self.glevel_fig.canvas.mpl_connect('button_press_event', self.on_glevel_plot_click)

        # Store plot index for zooming G-level plots
        self.glevel_plot_index = None

        # PSD plots tab
        self.psd_plots_frame = tk.Frame(self.notebook)
        self.notebook.add(self.psd_plots_frame, text="PSD Plots")

        # Canvas for plotting PSD
        self.psd_fig = plt.Figure(figsize=(14, 10))
        self.psd_canvas = FigureCanvasTkAgg(self.psd_fig, self.psd_plots_frame)
        self.psd_canvas.get_tk_widget().pack(side="top", fill="both", expand=True)

        # Toolbar for PSD plots
        self.psd_toolbar = NavigationToolbar2Tk(self.psd_canvas, self.psd_plots_frame)
        self.psd_toolbar.update()
        self.psd_toolbar.pack(side="bottom", fill="x")

        # Plot click event for PSD plots
        self.psd_fig.canvas.mpl_connect('button_press_event', self.on_psd_plot_click)

        # Store plot index for zooming PSD plots
        self.psd_plot_index = None

        # G-levels DataFrame
        self.data = None
        self.velocity_data = None  

        # Store image filenames
        self.image_filenames = []

        # Set initial state of save button
        self.save_button_state(False)

    def load_file(self, variable):
        file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv"), ("Excel files", "*.xlsx")])
        if file_path:
            variable.set(file_path)
            if variable == self.csv_file_path:
                self.data = self.read_data(file_path)
                messagebox.showinfo("Success", "Main Data loaded successfully.\nPath: {}".format(file_path))
            elif variable == self.velocity_csv_file_path:
                self.velocity_data = self.read_data(file_path)
                messagebox.showinfo("Success", "Velocity Data loaded successfully.\nPath: {}".format(file_path))

    def read_data(self, file_path):
        if file_path.endswith('.csv'):
            return pd.read_csv(file_path)
        elif file_path.endswith('.xlsx'):
            data = pd.read_excel(file_path)
            if 'Time' in data.columns:
                if data['Time'].max() > 100:
                    data['Time'] /= 1000  
            return data

    def plot_glevels(self):
        if self.data is not None and self.velocity_data is not None:
            num_channels = len(self.data.columns) - 1  
            num_plots = min(num_channels, 24)

            num_rows = (num_plots - 1) // 6 + 1
            num_cols = min(num_plots, 6)

            self.glevel_fig.clf()

            for i in range(num_plots):
                channel = self.data.columns[i + 1]  
                ax = self.glevel_fig.add_subplot(num_rows, num_cols, i + 1)

                ax.plot(self.data.iloc[:, 0], self.data.iloc[:, i + 1] / self.sensitivity.get(), label="G-levels", color='blue')

                ax2 = ax.twinx()
                ax2.plot(self.velocity_data.iloc[:, 0], self.velocity_data.iloc[:, 1], label="Velocity", color='orange')

                ax.set_title(channel, fontsize=8)
                ax.set_xlabel("Time (s)", fontsize=10)
                ax.set_ylabel("G-levels", fontsize=8)
                ax2.set_ylabel("Velocity", fontsize=8)

                ax.tick_params(axis='x', rotation=45, labelsize=8)
                ax.tick_params(axis='both', which='major', pad=2)

                ax.legend(loc='upper left')
                ax2.legend(loc='upper right')

            self.glevel_fig.subplots_adjust(hspace=1, wspace=0.5, top=0.95)

            self.glevel_canvas.draw()

            self.notebook.select(self.glevel_plots_frame)

            self.save_button_state(True)

    def plot_psd(self):
        if self.data is not None:
            num_channels = len(self.data.columns) - 1  
            num_plots = min(num_channels, 24)

            num_rows = (num_plots - 1) // 6 + 1
            num_cols = min(num_plots, 6)

            self.psd_fig.clf()

            for i in range(num_plots):
                channel = self.data.columns[i + 1]  
                ax = self.psd_fig.add_subplot(num_rows, num_cols, i + 1)

                f, p_s_d = welch(self.data.iloc[:, i + 1], fs=self.sampling_frequency.get())

                ax.semilogy(f, p_s_d)
                ax.set_title(channel, fontsize=8)
                ax.set_xlabel("Frequency (Hz)", fontsize=8)
                ax.set_ylabel("PSD", fontsize=8)

            self.psd_fig.subplots_adjust(hspace=1, wspace=0.5, top=0.95)

            self.psd_canvas.draw()

            self.notebook.select(self.psd_plots_frame)

            self.save_button_state(True)

    def save_button_state(self, state):
        if state:
            self.save_button.grid(row=8, column=0, columnspan=2)
        else:
            self.save_button.grid_remove()

    def save_images(self):
        if self.image_filenames:
            doc = Document()
            for filename in self.image_filenames:
                doc.add_picture(filename, width=Inches(6))
                doc.add_paragraph()
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                doc.save(save_path)
                messagebox.showinfo("Success", "Images saved successfully.\nPath: {}".format(save_path))
        else:
            messagebox.showinfo("No Images", "No images to save.")



    def on_glevel_plot_click(self, event):
        if event.inaxes and event.inaxes.get_figure() == self.glevel_fig:
            for i, ax in enumerate(self.glevel_fig.axes):
                if ax == event.inaxes:
                    self.glevel_plot_index = i
                    break

        if self.glevel_plot_index is not None and self.glevel_plot_index < len(self.data.columns) - 1:
            self.zoom_glevel_plot()


    def zoom_glevel_plot(self):
        if self.glevel_plot_index is not None:
            channel = self.data.columns[self.glevel_plot_index + 1]  

        fig, ax1 = plt.subplots(figsize=(8, 6))

        color = 'tab:blue'
        ax1.set_xlabel('Time (s)')
        ax1.set_ylabel('G-levels', color=color)
        ax1.plot(self.data.iloc[:, 0], self.data.iloc[:, self.glevel_plot_index + 1] / self.sensitivity.get(), color=color, label="G-levels")
        ax1.tick_params(axis='y', labelcolor=color)
        ax1.legend(loc='upper left')

        ax2 = ax1.twinx()
        color = 'tab:orange'
        ax2.set_ylabel('Velocity', color=color)
        ax2.plot(self.velocity_data.iloc[:, 0], self.velocity_data.iloc[:, 1], color=color, label="Velocity")
        ax2.tick_params(axis='y', labelcolor=color)
        ax2.legend(loc='upper right')

        plt.title(channel)
        image_filename = "glevel_{}.png".format(channel)
        plt.savefig(image_filename)
        plt.show()  
        plt.close()
        self.image_filenames.append(image_filename)


    def on_psd_plot_click(self, event):
        if event.inaxes and event.inaxes.get_figure() == self.psd_fig:
            for i, ax in enumerate(self.psd_fig.axes):
                if ax == event.inaxes:
                    self.psd_plot_index = i
                    break

            if self.psd_plot_index is not None:
                self.zoom_psd_plot()



    def zoom_psd_plot(self):
        if self.psd_plot_index is not None:
            plt.figure()
            channel = self.data.columns[self.psd_plot_index]  # Exclude time column

        # Calculate PSD using Welch method
        f, p_s_d = welch(self.data.iloc[:, self.psd_plot_index], fs=self.sampling_frequency.get())

        plt.semilogy(f, p_s_d, label="PSD")  # Add label here
        plt.title(channel)
        plt.xlabel("Frequency (Hz)")
        plt.ylabel("PSD")
        plt.legend()  # Add legend here
        image_filename = "psd_{}.png".format(channel)
        plt.savefig(image_filename)
        plt.close()
        self.image_filenames.append(image_filename)






if __name__ == "__main__":
    root = tk.Tk()
    app = VibrationAnalyzer(root)
    root.mainloop()
